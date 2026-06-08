import os
from typing import Dict, Any, Optional
from datetime import datetime

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Word imports
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

class ReportingService:
    @classmethod
    def generate_pdf_report(
        self, 
        output_path: str, 
        report_name: str, 
        dataset_name: str, 
        eda_data: Dict[str, Any], 
        anomalies_data: Optional[Dict[str, Any]] = None,
        ai_insights: Optional[str] = None
    ) -> str:
        """Generates a professional PDF report containing the executive summary, EDA, and anomaly data."""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54, leftMargin=54,
            topMargin=54, bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Define Custom Color Palette (Premium Indigo Theme)
        primary_color = colors.HexColor("#3b82f6")  # Sleek Blue
        secondary_color = colors.HexColor("#1e3a8a")  # Dark Blue
        dark_neutral = colors.HexColor("#1f2937")  # Charcoal
        light_neutral = colors.HexColor("#f3f4f6")  # Cool Grey
        accent_color = colors.HexColor("#f59e0b")  # Amber Accent
        
        # Modify existing styles to avoid crashes
        styles['Normal'].textColor = dark_neutral
        styles['Normal'].fontSize = 10
        styles['Normal'].leading = 14
        
        # Add new custom styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=28,
            leading=34,
            textColor=secondary_color,
            spaceAfter=15
        )
        
        subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=styles['Normal'],
            fontSize=13,
            leading=18,
            textColor=colors.HexColor("#6b7280"),
            spaceAfter=30
        )
        
        h1_style = ParagraphStyle(
            'Heading1_Custom',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=secondary_color,
            spaceBefore=15,
            spaceAfter=10,
            keepWithNext=True
        )
        
        h2_style = ParagraphStyle(
            'Heading2_Custom',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=13,
            leading=16,
            textColor=primary_color,
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True
        )
        
        bullet_style = ParagraphStyle(
            'Bullet_Custom',
            parent=styles['Normal'],
            leftIndent=15,
            spaceAfter=4
        )

        story = []
        
        # ================= PAGE 1: COVER PAGE =================
        story.append(Spacer(1, 1.5 * inch))
        story.append(Paragraph(report_name, title_style))
        story.append(Paragraph(f"Intelligent Data Analysis Report • Dataset: {dataset_name}", subtitle_style))
        
        # Decorative line
        d_table = Table([[""]], colWidths=[504], rowHeights=[4])
        d_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), primary_color),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(d_table)
        story.append(Spacer(1, 0.4 * inch))
        
        # Metadata block
        meta_data = [
            [Paragraph("<b>Author:</b> AI Analyst Agent", styles['Normal']), 
             Paragraph(f"<b>Date:</b> {datetime.now().strftime('%B %d, %Y')}", styles['Normal'])],
            [Paragraph("<b>Status:</b> Completed & Verified", styles['Normal']), 
             Paragraph("<b>Security:</b> Restricted / Internal", styles['Normal'])]
        ]
        meta_table = Table(meta_data, colWidths=[252, 252])
        meta_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), light_neutral),
            ('PADDING', (0,0), (-1,-1), 10),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(meta_table)
        story.append(PageBreak())
        
        # ================= PAGE 2: EXECUTIVE SUMMARY =================
        story.append(Paragraph("Executive Summary", h1_style))
        story.append(Spacer(1, 0.1 * inch))
        
        if ai_insights:
            # Parse AI summary markdown and append it
            paragraphs = ai_insights.split("\n\n")
            for p in paragraphs:
                p = p.strip()
                if p.startswith("### "):
                    story.append(Paragraph(p.replace("### ", ""), h2_style))
                elif p.startswith("## "):
                    story.append(Paragraph(p.replace("## ", ""), h1_style))
                elif p.startswith("- ") or p.startswith("* "):
                    lines = p.split("\n")
                    for line in lines:
                        cleaned_line = line.strip().lstrip("-* ").strip()
                        story.append(Paragraph(f"• {cleaned_line}", bullet_style))
                elif p:
                    story.append(Paragraph(p, styles['Normal']))
                    story.append(Spacer(1, 6))
        else:
            story.append(Paragraph("No AI insights generated for this dataset yet.", styles['Normal']))
            
        story.append(Spacer(1, 0.2 * inch))
        
        # ================= SECTION 3: DATASET OVERVIEW =================
        story.append(Paragraph("Dataset Overview", h1_style))
        overview = eda_data.get("overview", {})
        
        overview_data = [
            ["Metric", "Value"],
            ["Total Rows", f"{overview.get('row_count', 0):,}"],
            ["Total Columns", str(overview.get('col_count', 0))],
            ["Duplicate Rows", f"{overview.get('duplicate_count', 0):,}"],
            ["Missing Cells", f"{overview.get('missing_count', 0):,} ({overview.get('missing_percentage', 0):.2f}%)"]
        ]
        
        overview_table = Table(overview_data, colWidths=[200, 304])
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), secondary_color),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 10),
            ('BACKGROUND', (0,1), (-1,-1), light_neutral),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e5e7eb")),
            ('PADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(overview_table)
        story.append(Spacer(1, 0.2 * inch))
        
        # ================= SECTION 4: DATA QUALITY & COLUMNS =================
        story.append(Paragraph("Column Inventory & Quality Profile", h1_style))
        
        columns_list = eda_data.get("columns", [])
        col_headers = ["Name", "Type", "Nulls (%)", "Unique Values"]
        col_table_data = [col_headers]
        
        for c in columns_list[:12]:  # Display top 12 columns to fit properly
            col_table_data.append([
                c["name"][:25],
                c["type"],
                f"{c['null_count']} ({c['null_percentage']:.1f}%)",
                f"{c['unique_count']:,}"
            ])
            
        col_table = Table(col_table_data, colWidths=[160, 114, 130, 100])
        col_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), primary_color),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, light_neutral]),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e5e7eb")),
            ('PADDING', (0,0), (-1,-1), 5),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(col_table)
        
        if len(columns_list) > 12:
            story.append(Paragraph(f"<i>* Showing first 12 of {len(columns_list)} columns. View system dashboard for full list.</i>", subtitle_style))
            
        # ================= SECTION 5: ANOMALY REPORT =================
        if anomalies_data and anomalies_data.get("success"):
            story.append(PageBreak())
            story.append(Paragraph("Statistical Anomaly Analysis", h1_style))
            
            summary = anomalies_data.get("summary", {})
            story.append(Paragraph(
                f"We scanned the numerical columns for statistical outliers. A total of <b>{summary.get('total_detected', 0)}</b> anomalies "
                f"were detected (constituting {summary.get('percentage_of_data', 0):.2f}% of overall dataset records). "
                f"Breakdown of severity: <b>{summary.get('high_severity_count', 0)} High</b>, <b>{summary.get('medium_severity_count', 0)} Medium</b>.",
                styles['Normal']
            ))
            story.append(Spacer(1, 0.15 * inch))
            
            anom_list = anomalies_data.get("anomalies", [])
            if anom_list:
                anom_headers = ["Col", "Val", "Severity", "Confidence", "Impact"]
                anom_table_data = [anom_headers]
                
                for a in anom_list[:8]:  # Top 8 anomalies
                    col_name = a.get("column", "Multi") if "column" in a else "Multi-col"
                    val = f"{a.get('value', 0):.2f}" if "value" in a and a.get("value") is not None else "N/A"
                    anom_table_data.append([
                        col_name[:15],
                        val,
                        a.get("severity", "low").upper(),
                        f"{a.get('confidence', 0)*100:.0f}%",
                        a.get("impact", "")[:45] + "..."
                    ])
                    
                anom_table = Table(anom_table_data, colWidths=[110, 80, 80, 80, 154])
                anom_table.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#dc2626")),  # Red Header
                    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, light_neutral]),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#e5e7eb")),
                    ('PADDING', (0,0), (-1,-1), 5),
                    ('FONTSIZE', (0,0), (-1,-1), 9),
                    ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ]))
                story.append(anom_table)
                story.append(Spacer(1, 0.1 * inch))
                
        # Build Document
        doc.build(story)
        return output_path

    @classmethod
    def generate_word_report(
        self, 
        output_path: str, 
        report_name: str, 
        dataset_name: str, 
        eda_data: Dict[str, Any], 
        anomalies_data: Optional[Dict[str, Any]] = None,
        ai_insights: Optional[str] = None
    ) -> str:
        """Generates a professional DOCX file."""
        doc = docx.Document()
        
        # Configure margins
        sections = doc.sections
        for s in sections:
            s.top_margin = Inches(1)
            s.bottom_margin = Inches(1)
            s.left_margin = Inches(1)
            s.right_margin = Inches(1)
            
        # Palette references (RGB Color definitions)
        c_primary = RGBColor(0x3B, 0x82, 0xF6)   # 59, 130, 246
        c_secondary = RGBColor(0x1E, 0x3A, 0x8A) # 30, 58, 138
        
        # Helpers for styling text
        def add_heading_1(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = c_secondary
            return p

        def add_heading_2(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = c_primary
            return p

        # Cover details
        p_title = doc.add_paragraph()
        p_title.paragraph_format.space_before = Pt(100)
        p_title.paragraph_format.space_after = Pt(12)
        run_title = p_title.add_run(report_name)
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(28)
        run_title.font.bold = True
        run_title.font.color.rgb = c_secondary
        
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(240)
        run_sub = p_sub.add_run(f"Intelligent Data Analysis Report\nDataset: {dataset_name}\nCreated: {datetime.now().strftime('%B %d, %Y')}")
        run_sub.font.name = 'Arial'
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        
        doc.add_page_break()
        
        # Executive Summary
        add_heading_1("Executive Summary")
        if ai_insights:
            for p_text in ai_insights.split("\n\n"):
                p_text = p_text.strip()
                if p_text.startswith("### "):
                    add_heading_2(p_text.replace("### ", ""))
                elif p_text.startswith("## "):
                    add_heading_1(p_text.replace("## ", ""))
                elif p_text.startswith("- ") or p_text.startswith("* "):
                    for line in p_text.split("\n"):
                        cleaned = line.strip().lstrip("-* ").strip()
                        doc.add_paragraph(cleaned, style='List Bullet')
                elif p_text:
                    p = doc.add_paragraph(p_text)
                    p.style.font.name = 'Calibri'
                    p.style.font.size = Pt(11)
        else:
            doc.add_paragraph("No AI summaries generated.")
            
        doc.add_page_break()
        
        # Overview
        add_heading_1("Dataset Overview")
        overview = eda_data.get("overview", {})
        
        # Table mapping
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        overview_metrics = [
            ("Total Rows", f"{overview.get('row_count', 0):,}"),
            ("Total Columns", str(overview.get('col_count', 0))),
            ("Duplicate Rows", f"{overview.get('duplicate_count', 0):,}"),
            ("Missing Cells", f"{overview.get('missing_count', 0):,} ({overview.get('missing_percentage', 0):.2f}%)")
        ]
        
        for metric, val in overview_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = val
            
        doc.save(output_path)
        return output_path
